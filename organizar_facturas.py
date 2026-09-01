#!/usr/bin/env python3
"""Crea un Word con facturas PDF ordenadas por la fecha de su nombre.

Ejemplos de nombres admitidos:
    20-08-2026 11410708.pdf       -> fecha 20/08/2026, contraseña 11410708
    21-08-2026.pdf                -> fecha 21/08/2026, sin contraseña
    2026-08-21_11410708.pdf       -> fecha 21/08/2026, contraseña 11410708
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import re
import sys
import tempfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm


DATE_PATTERNS = (
    re.compile(r"(?P<day>\d{2})[-_.](?P<month>\d{2})[-_.](?P<year>\d{4})"),
    re.compile(r"(?P<year>\d{4})[-_.](?P<month>\d{2})[-_.](?P<day>\d{2})"),
)


@dataclass(frozen=True)
class Invoice:
    path: Path
    date: datetime
    password: str | None


def parse_invoice_name(path: Path) -> Invoice:
    """Obtiene fecha y contraseña desde el nombre, sin incluir la extensión."""
    stem = path.stem.strip()
    for pattern in DATE_PATTERNS:
        match = pattern.search(stem)
        if not match:
            continue
        try:
            date = datetime(**{key: int(value) for key, value in match.groupdict().items()})
        except ValueError:
            continue

        # Todo lo que siga a la fecha, separado por espacios, guiones o guiones bajos,
        # es la contraseña. Se elimina el sufijo que Windows añade a copias: "(1)".
        remainder = stem[match.end() :].strip(" _-")
        remainder = re.sub(r"\s*\(\d+\)\s*$", "", remainder).strip()
        return Invoice(path=path, date=date, password=remainder or None)

    raise ValueError("No se encontró una fecha válida (DD-MM-AAAA o AAAA-MM-DD).")


def discover_invoices(input_dir: Path) -> tuple[list[Invoice], list[tuple[Path, str]]]:
    invoices: list[Invoice] = []
    errors: list[tuple[Path, str]] = []
    for path in sorted(input_dir.rglob("*"), key=lambda item: item.name.casefold()):
        if not path.is_file() or path.suffix.casefold() != ".pdf":
            continue
        try:
            invoices.append(parse_invoice_name(path))
        except ValueError as error:
            errors.append((path, str(error)))
    return sorted(invoices, key=lambda item: (item.date, item.path.name.casefold())), errors


def image_name(invoice: Invoice, page_number: int) -> str:
    digest = hashlib.sha256(str(invoice.path.resolve()).encode()).hexdigest()[:12]
    return f"{invoice.date:%Y%m%d}_{digest}_{page_number:03d}.png"


def render_pages(invoice: Invoice, image_dir: Path) -> list[Path]:
    """Abre el PDF, aplica la contraseña del nombre si hace falta y lo rasteriza."""
    document = fitz.open(invoice.path)
    try:
        if document.needs_pass:
            password = invoice.password or ""
            if not document.authenticate(password):
                if invoice.password:
                    raise RuntimeError("La contraseña indicada en el nombre no abrió el PDF.")
                raise RuntimeError("El PDF está protegido y el nombre no contiene contraseña.")

        images: list[Path] = []
        for number, page in enumerate(document, start=1):
            # 144 dpi conserva la legibilidad y evita Word excesivamente pesado.
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image_path = image_dir / image_name(invoice, number)
            pixmap.save(image_path)
            images.append(image_path)
        return images
    finally:
        document.close()


def add_invoice(document: Document, invoice: Invoice, number: int, image_dir: Path) -> None:
    title = document.add_heading(f"FACTURA {number}", level=2)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details = document.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.add_run(f"Fecha: {invoice.date:%d/%m/%Y}\n").bold = True
    details.add_run(f"Archivo: {invoice.path.name}")

    for image in render_pages(invoice, image_dir):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(image), width=Cm(18.5))


def write_errors(errors: list[tuple[Path, str]], output_dir: Path) -> None:
    if not errors:
        return
    report = output_dir / "facturas_con_error.csv"
    with report.open("w", newline="", encoding="utf-8-sig") as file:
        writer = csv.writer(file)
        writer.writerow(("archivo", "motivo"))
        writer.writerows((str(path), reason) for path, reason in errors)


def create_word(invoices: list[Invoice], output_path: Path, errors: list[tuple[Path, str]]) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)

    heading = document.add_heading("FACTURAS ORDENADAS POR FECHA", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Facturas incluidas: {len(invoices)}")
    document.add_paragraph("Orden: de la fecha más antigua a la más reciente.")

    with tempfile.TemporaryDirectory(prefix="facturas_word_") as temporary:
        image_dir = Path(temporary)
        included = 0
        for number, invoice in enumerate(invoices, start=1):
            try:
                document.add_page_break()
                add_invoice(document, invoice, number, image_dir)
                included += 1
                print(f"Incluida: {invoice.date:%d/%m/%Y} — {invoice.path.name}")
            except Exception as error:  # Se conserva el resto de facturas si una falla.
                errors.append((invoice.path, str(error)))
                print(f"ERROR: {invoice.path.name}: {error}", file=sys.stderr)

    if included == 0:
        raise RuntimeError("No se pudo incluir ninguna factura; no se creó el Word.")
    document.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", nargs="?", default="Facturas", type=Path, help="Carpeta con los PDF (predeterminada: Facturas).")
    parser.add_argument("--output", default=Path("Facturas_ordenadas.docx"), type=Path, help="Ruta del Word a crear.")
    arguments = parser.parse_args()

    if not arguments.input.is_dir():
        print(f"No existe la carpeta de entrada: {arguments.input}", file=sys.stderr)
        return 2

    invoices, errors = discover_invoices(arguments.input)
    if not invoices:
        print("No se encontraron PDF con una fecha válida.", file=sys.stderr)
        write_errors(errors, arguments.output.parent)
        return 1

    arguments.output.parent.mkdir(parents=True, exist_ok=True)
    create_word(invoices, arguments.output, errors)
    write_errors(errors, arguments.output.parent)
    print(f"\nWord creado: {arguments.output.resolve()}")
    if errors:
        print(f"Advertencia: {len(errors)} archivo(s) no se incluyeron. Revise facturas_con_error.csv.")
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
